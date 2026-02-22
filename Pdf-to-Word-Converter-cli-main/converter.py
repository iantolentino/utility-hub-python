import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import tempfile
import traceback
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Inches
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
from pdf2docx import Converter  

class PDFtoDocxConverter:
    """Initialized the file, and setting up paths for the converted files"""
    def __init__(self, root):
        self.root = root
        self.root.title("PDF to DOCX Converter")
        self.root.geometry("800x600")
        
        # Initialize variables
        self.input_file = None
        self.output_path = tk.StringVar(value=str(Path.home() / "Documents" / "Converted_Files"))
        self.setup_ui()

    """Contains UI setup and layout"""
    def setup_ui(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        title_label = ttk.Label(main_frame, text="PDF to DOCX Converter", font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))

        file_frame = ttk.LabelFrame(main_frame, text="File Selection", padding="10")
        file_frame.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)

        # Single file input
        self.single_file_frame = ttk.Frame(file_frame)
        self.single_file_frame.grid(row=0, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        self.single_file_frame.columnconfigure(0, weight=1)

        self.single_file_path = tk.StringVar()
        single_file_entry = ttk.Entry(self.single_file_frame, textvariable=self.single_file_path, state='readonly')
        single_file_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))

        ttk.Button(self.single_file_frame, text="Browse File", command=self.browse_single_file).grid(row=0, column=1)

        output_frame = ttk.LabelFrame(main_frame, text="Output Location", padding="10")
        output_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        output_frame.columnconfigure(0, weight=1)

        output_entry = ttk.Entry(output_frame, textvariable=self.output_path)
        output_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
        
        ttk.Button(output_frame, text="Browse", command=self.browse_output_location).grid(row=0, column=1)

        convert_btn = ttk.Button(main_frame, text="Start Conversion", command=self.start_conversion)
        convert_btn.grid(row=3, column=0, columnspan=3, pady=20)

        self.status_label = ttk.Label(main_frame, text="Ready to convert")
        self.status_label.grid(row=4, column=0, columnspan=3)
        
    def browse_single_file(self):
        """Validator if the file is pdf based on the selection of the user"""
        file_path = filedialog.askopenfilename(
            title="Select PDF File",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        if file_path:
            self.single_file_path.set(file_path)
    
    def browse_output_location(self):
        """Where the user wants the file output"""
        folder = filedialog.askdirectory(title="Select Output Location")
        if folder:
            self.output_path.set(folder)
    
    def start_conversion(self):
        """Conversion initial setup that has a required file to be uploaded, and a valid output location"""
        input_file = self.single_file_path.get()
        output_folder = self.output_path.get()
        
        if not input_file or not os.path.exists(input_file):
            messagebox.showwarning("Warning", "Please select a valid PDF file to convert.")
            return
        
        if not output_folder or not os.path.exists(output_folder):
            messagebox.showwarning("Warning", "Please select a valid output location.")
            return
        
        self.status_label.config(text="Converting...")

        # Perform conversion
        output_file = os.path.join(output_folder, os.path.basename(input_file).replace(".pdf", ".docx"))
        
        try:
            # Use pdf2docx if available
            success = self.convert_with_pdf2docx(input_file, output_file)
            
            if not success:
                # Fallback to image-based or text-based methods
                success = self.convert_with_fitz(input_file, output_file)
            
            if success:
                messagebox.showinfo("Success", f"Conversion successful! File saved at:\n{output_file}")
            else:
                messagebox.showerror("Error", "Conversion failed.")
            
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred during conversion:\n{str(e)}")
            print(f"Error: {traceback.format_exc()}")
        
        self.status_label.config(text="Ready to convert")
    
    def convert_with_pdf2docx(self, input_file, output_file):
        """Use pdf2docx library for better layout preservation."""
        try:
            cv = Converter(input_file)
            cv.convert(output_file, start=0, end=None)
            cv.close()
            print(f"pdf2docx conversion successful: {output_file}")
            return True
        except Exception as e:
            print(f"pdf2docx conversion failed: {e}")
            return False

    def convert_with_fitz(self, input_file, output_file):
        """Use PyMuPDF (fitz) for text extraction and image embedding."""
        try:
            doc = Document()
            pdf_document = fitz.open(input_file)

            # Add title
            doc.add_heading(os.path.basename(input_file), 0)

            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text = page.get_text("dict")

                if text:
                    for block in text["blocks"]:
                        if block["type"] == 0:  # Text block
                            for line in block["lines"]:
                                paragraph = " ".join([span['text'] for span in line["spans"]])
                                if paragraph.strip():
                                    doc.add_paragraph(paragraph)

                # Add image if page contains image content
                pix = page.get_pixmap()
                img_path = os.path.join(tempfile.gettempdir(), f"page_{page_num+1}.png")
                pix.save(img_path)
                doc.add_picture(img_path, width=Inches(7.5))
                
                if page_num < len(pdf_document) - 1:
                    doc.add_page_break()

            pdf_document.close()
            doc.save(output_file)
            print(f"fitz conversion successful: {output_file}")
            return True

        except Exception as e:
            print(f"fitz conversion failed: {e}")
            return False

def main():
    root = tk.Tk()
    app = PDFtoDocxConverter(root)
    root.mainloop()

if __name__ == "__main__":
    main()
