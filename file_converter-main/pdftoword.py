import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
from pathlib import Path
import threading
import sys
import traceback
import subprocess
import shutil
import tempfile

class UniversalPDFConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("Universal PDF Converter")
        self.root.geometry("800x600")
        self.root.resizable(True, True)
        
        # Check available libraries
        self.libraries = self.check_available_libraries()
        
        # Variables
        self.input_files = []
        self.output_format = tk.StringVar(value="docx")
        self.conversion_mode = tk.StringVar(value="single")
        
        self.setup_ui()
        
    def check_available_libraries(self):
        """Check which conversion libraries are available"""
        libraries = {
            'pdf2image': False,
            'pypdf': False,
            'fitz': False,
            'python_docx': False,
        }
        
        try:
            from pdf2image import convert_from_path
            libraries['pdf2image'] = True
        except ImportError:
            pass
            
        try:
            import PyPDF2
            libraries['pypdf'] = True
        except ImportError:
            pass
            
        try:
            import fitz  # PyMuPDF
            libraries['fitz'] = True
        except ImportError:
            pass
            
        try:
            from docx import Document
            from docx.shared import Inches
            libraries['python_docx'] = True
        except ImportError:
            pass
            
        return libraries
    
    def setup_ui(self):
        # Main container
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title
        title_label = ttk.Label(main_frame, text="Universal PDF Converter", 
                               font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        # Library Status
        status_frame = ttk.LabelFrame(main_frame, text="Library Status", padding="5")
        status_frame.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        status_text = "Available: "
        available_libs = [lib for lib, available in self.libraries.items() if available]
        status_text += ", ".join(available_libs) if available_libs else "None"
        
        status_label = ttk.Label(status_frame, text=status_text, 
                               foreground="green" if available_libs else "red")
        status_label.pack()
        
        # Install missing libraries button
        if not all(self.libraries.values()):
            ttk.Button(status_frame, text="Install Missing Libraries", 
                      command=self.install_all_missing_libraries).pack(pady=5)
        
        # Format Selection
        format_frame = ttk.LabelFrame(main_frame, text="Conversion Format", padding="10")
        format_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        formats = [
            ("Word Document (.docx)", "docx"),
            ("Images (PNG)", "png"),
            ("Images (JPG)", "jpg"),
            ("Text File (.txt)", "txt"),
            ("PDF (Copy)", "pdf")
        ]
        
        for i, (text, value) in enumerate(formats):
            ttk.Radiobutton(format_frame, text=text, variable=self.output_format, 
                           value=value).grid(row=0, column=i, sticky=tk.W, padx=5)
        
        # Mode Selection
        mode_frame = ttk.LabelFrame(main_frame, text="Conversion Mode", padding="10")
        mode_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        ttk.Radiobutton(mode_frame, text="Single File", variable=self.conversion_mode, 
                       value="single").grid(row=0, column=0, sticky=tk.W, padx=5)
        ttk.Radiobutton(mode_frame, text="Batch Files", variable=self.conversion_mode, 
                       value="batch").grid(row=0, column=1, sticky=tk.W, padx=5)
        ttk.Radiobutton(mode_frame, text="Folder", variable=self.conversion_mode, 
                       value="folder").grid(row=0, column=2, sticky=tk.W, padx=5)
        
        # File Selection
        file_frame = ttk.LabelFrame(main_frame, text="File Selection", padding="10")
        file_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        file_frame.columnconfigure(0, weight=1)
        
        # Single file input
        self.single_file_frame = ttk.Frame(file_frame)
        self.single_file_frame.grid(row=0, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        self.single_file_frame.columnconfigure(0, weight=1)
        
        self.single_file_path = tk.StringVar()
        single_file_entry = ttk.Entry(self.single_file_frame, textvariable=self.single_file_path, state='readonly')
        single_file_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
        
        ttk.Button(self.single_file_frame, text="Browse File", 
                  command=self.browse_single_file).grid(row=0, column=1)
        
        # Batch files input
        self.batch_files_frame = ttk.Frame(file_frame)
        self.batch_files_frame.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        self.batch_files_frame.columnconfigure(0, weight=1)
        
        self.batch_files_listbox = tk.Listbox(self.batch_files_frame, height=6)
        self.batch_files_listbox.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        batch_btn_frame = ttk.Frame(self.batch_files_frame)
        batch_btn_frame.grid(row=0, column=2, sticky=(tk.N, tk.S), padx=5)
        
        ttk.Button(batch_btn_frame, text="Add Files", 
                  command=self.add_batch_files).pack(fill=tk.X, pady=2)
        ttk.Button(batch_btn_frame, text="Remove", 
                  command=self.remove_batch_file).pack(fill=tk.X, pady=2)
        ttk.Button(batch_btn_frame, text="Clear All", 
                  command=self.clear_batch_files).pack(fill=tk.X, pady=2)
        
        # Folder input
        self.folder_frame = ttk.Frame(file_frame)
        self.folder_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        self.folder_frame.columnconfigure(0, weight=1)
        
        self.folder_path = tk.StringVar()
        folder_entry = ttk.Entry(self.folder_frame, textvariable=self.folder_path, state='readonly')
        folder_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
        
        ttk.Button(self.folder_frame, text="Browse Folder", 
                  command=self.browse_folder).grid(row=0, column=1)
        
        # Show/hide frames based on mode
        self.update_mode_display()
        self.conversion_mode.trace('w', self.on_mode_change)
        
        # Output Location
        output_frame = ttk.LabelFrame(main_frame, text="Output Location", padding="10")
        output_frame.grid(row=5, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        output_frame.columnconfigure(0, weight=1)
        
        self.output_path = tk.StringVar(value=str(Path.home() / "Documents" / "Converted_Files"))
        output_entry = ttk.Entry(output_frame, textvariable=self.output_path)
        output_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
        
        ttk.Button(output_frame, text="Browse", 
                  command=self.browse_output_location).grid(row=0, column=1)
        
        # Conversion Method for DOCX
        method_frame = ttk.LabelFrame(main_frame, text="DOCX Conversion Method", padding="10")
        method_frame.grid(row=6, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        self.docx_method = tk.StringVar(value="image_based")
        
        ttk.Radiobutton(method_frame, text="Image-based (Best for scanned PDFs)", 
                       variable=self.docx_method, value="image_based").grid(row=0, column=0, sticky=tk.W)
        ttk.Radiobutton(method_frame, text="Text extraction (For text-based PDFs)", 
                       variable=self.docx_method, value="text_based").grid(row=0, column=1, sticky=tk.W)
        
        # Options Frame
        options_frame = ttk.LabelFrame(main_frame, text="Conversion Options", padding="10")
        options_frame.grid(row=7, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        self.high_quality = tk.BooleanVar(value=True)
        self.ocr_enabled = tk.BooleanVar(value=False)
        
        ttk.Checkbutton(options_frame, text="High quality (300 DPI)", 
                       variable=self.high_quality).grid(row=0, column=0, sticky=tk.W)
        ttk.Checkbutton(options_frame, text="Enable OCR (Experimental)", 
                       variable=self.ocr_enabled).grid(row=0, column=1, sticky=tk.W)
        
        # Convert Button
        self.convert_btn = ttk.Button(main_frame, text="Start Conversion", 
                                     command=self.start_conversion)
        self.convert_btn.grid(row=8, column=0, columnspan=3, pady=20)
        
        # Progress
        self.progress_frame = ttk.Frame(main_frame)
        self.progress_frame.grid(row=9, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        self.progress_bar = ttk.Progressbar(self.progress_frame, mode='determinate')
        self.progress_bar.pack(fill=tk.X, expand=True)
        
        self.status_label = ttk.Label(main_frame, text="Ready to convert")
        self.status_label.grid(row=10, column=0, columnspan=3)
        
        # Results
        results_frame = ttk.LabelFrame(main_frame, text="Conversion Results", padding="10")
        results_frame.grid(row=11, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        results_frame.columnconfigure(0, weight=1)
        results_frame.rowconfigure(0, weight=1)
        main_frame.rowconfigure(11, weight=1)
        
        self.results_text = tk.Text(results_frame, height=8, wrap=tk.WORD)
        self.results_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        scrollbar = ttk.Scrollbar(results_frame, orient="vertical", command=self.results_text.yview)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.results_text.configure(yscrollcommand=scrollbar.set)
        
        # Web conversion suggestion
        suggestion_label = ttk.Label(main_frame, 
                                   text="💡 Tip: For complex PDFs with tables/layout, consider using online converters like SmallPDF, iLovePDF, or Adobe Online",
                                   foreground="blue", font=("Arial", 9))
        suggestion_label.grid(row=12, column=0, columnspan=3, pady=10)
        
    def install_all_missing_libraries(self):
        """Install all missing libraries"""
        missing_libs = [lib for lib, available in self.libraries.items() if not available]
        
        if not missing_libs:
            messagebox.showinfo("Info", "All required libraries are already installed!")
            return
            
        result = messagebox.askyesno(
            "Install Missing Libraries",
            f"The following libraries are missing:\n{', '.join(missing_libs)}\n\n"
            f"Would you like to install them now?"
        )
        
        if result:
            try:
                self.status_label.config(text="Installing missing libraries...")
                
                for lib in missing_libs:
                    if lib == 'pdf2image':
                        subprocess.check_call([sys.executable, "-m", "pip", "install", "pdf2image", "pillow"])
                    elif lib == 'python_docx':
                        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
                    else:
                        subprocess.check_call([sys.executable, "-m", "pip", "install", lib])
                
                messagebox.showinfo("Success", "All libraries installed successfully!\nPlease restart the application.")
                self.status_label.config(text="Libraries installed - Please restart")
                
            except Exception as e:
                messagebox.showerror("Installation Failed", f"Failed to install libraries:\n{str(e)}")
                self.status_label.config(text="Installation failed")
        
    def on_mode_change(self, *args):
        self.update_mode_display()
        
    def update_mode_display(self):
        mode = self.conversion_mode.get()
        
        # Hide all frames first
        self.single_file_frame.grid_remove()
        self.batch_files_frame.grid_remove()
        self.folder_frame.grid_remove()
        
        # Show selected frame
        if mode == "single":
            self.single_file_frame.grid()
        elif mode == "batch":
            self.batch_files_frame.grid()
        elif mode == "folder":
            self.folder_frame.grid()
    
    def browse_single_file(self):
        file_path = filedialog.askopenfilename(
            title="Select PDF File",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        if file_path:
            self.single_file_path.set(file_path)
    
    def add_batch_files(self):
        files = filedialog.askopenfilenames(
            title="Select PDF Files",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        for file in files:
            if file not in self.input_files:
                self.input_files.append(file)
                self.batch_files_listbox.insert(tk.END, os.path.basename(file))
    
    def remove_batch_file(self):
        selection = self.batch_files_listbox.curselection()
        if selection:
            index = selection[0]
            self.input_files.pop(index)
            self.batch_files_listbox.delete(index)
    
    def clear_batch_files(self):
        self.input_files.clear()
        self.batch_files_listbox.delete(0, tk.END)
    
    def browse_folder(self):
        folder = filedialog.askdirectory(title="Select Folder with PDF Files")
        if folder:
            self.folder_path.set(folder)
    
    def browse_output_location(self):
        folder = filedialog.askdirectory(title="Select Output Location")
        if folder:
            self.output_path.set(folder)
    
    def get_files_to_convert(self):
        mode = self.conversion_mode.get()
        files = []
        
        if mode == "single":
            if self.single_file_path.get():
                files.append(self.single_file_path.get())
        elif mode == "batch":
            files = self.input_files.copy()
        elif mode == "folder":
            folder = self.folder_path.get()
            if folder and os.path.exists(folder):
                for file in Path(folder).glob("*.pdf"):
                    files.append(str(file))
        
        return files
    
    def start_conversion(self):
        # Debug information
        print("=== DEBUG INFO ===")
        print(f"Available libraries: {self.libraries}")
        print(f"Output format: {self.output_format.get()}")
        print(f"DOCX method: {self.docx_method.get()}")
        print(f"Python version: {sys.version}")
        print("==================")
        
        files = self.get_files_to_convert()
        
        if not files:
            messagebox.showwarning("Warning", "Please select files to convert.")
            return
        
        if not self.output_path.get():
            messagebox.showwarning("Warning", "Please select an output location.")
            return
        
        # Check if required libraries are available
        target_format = self.output_format.get()
        if target_format == "docx" and not (self.libraries['fitz'] and self.libraries['python_docx']):
            messagebox.showerror("Missing Libraries", "PyMuPDF and python-docx are required for DOCX conversion.")
            return
        elif target_format in ["png", "jpg"] and not self.libraries['pdf2image']:
            self.install_missing_library("pdf2image")
            return
        
        # Disable convert button during conversion
        self.convert_btn.config(state='disabled')
        self.progress_bar.config(value=0, maximum=len(files))
        
        # Start conversion in thread
        thread = threading.Thread(target=self.convert_files, args=(files,))
        thread.daemon = True
        thread.start()
    
    def convert_files(self, files):
        successful = 0
        failed = 0
        failed_files = []
        
        output_dir = self.output_path.get()
        os.makedirs(output_dir, exist_ok=True)
        target_format = self.output_format.get()
        
        for i, file_path in enumerate(files):
            try:
                self.update_status(f"Converting {i+1}/{len(files)}: {os.path.basename(file_path)}")
                self.update_progress(i + 1)
                
                file_name = os.path.splitext(os.path.basename(file_path))[0]
                output_path = os.path.join(output_dir, f"{file_name}.{target_format}")
                
                # Handle duplicate files
                counter = 1
                while os.path.exists(output_path):
                    output_path = os.path.join(output_dir, f"{file_name}_{counter}.{target_format}")
                    counter += 1
                
                # Perform conversion based on target format
                success = False
                if target_format == "docx":
                    if self.docx_method.get() == "image_based":
                        success = self.convert_to_docx_image_based(file_path, output_path)
                    else:
                        success = self.convert_to_docx_text_based(file_path, output_path)
                elif target_format in ["png", "jpg"]:
                    success = self.convert_to_image(file_path, output_path, target_format)
                elif target_format == "txt":
                    success = self.convert_to_text(file_path, output_path)
                elif target_format == "pdf":
                    success = self.copy_pdf(file_path, output_path)
                else:
                    success = False
                
                if success:
                    successful += 1
                    self.add_result(f"✓ {os.path.basename(file_path)} → {os.path.basename(output_path)}")
                else:
                    failed += 1
                    failed_files.append(os.path.basename(file_path))
                    self.add_result(f"✗ {os.path.basename(file_path)} - Conversion failed")
                    
            except Exception as e:
                failed += 1
                failed_files.append(os.path.basename(file_path))
                error_details = f"✗ {os.path.basename(file_path)} - Error: {str(e)}"
                self.add_result(error_details)
                print(f"Conversion error details: {traceback.format_exc()}")
        
        # Final update
        self.root.after(0, self.conversion_complete, successful, failed, failed_files)
    
    def convert_to_docx_image_based(self, pdf_path, output_path):
        """Convert PDF to Word by embedding pages as images - BEST FOR SCANNED PDFs"""
        try:
            if not self.libraries['fitz'] or not self.libraries['python_docx']:
                return False
                
            import fitz
            from docx import Document
            from docx.shared import Inches
            import tempfile
            import os
            
            print(f"Using image-based conversion for {pdf_path}")
            
            pdf_document = fitz.open(pdf_path)
            doc = Document()
            
            # Set DPI based on quality setting
            dpi = 300 if self.high_quality.get() else 150
            
            # Create temp directory for images
            with tempfile.TemporaryDirectory() as temp_dir:
                # Convert each page to image
                for page_num in range(len(pdf_document)):
                    page = pdf_document.load_page(page_num)
                    
                    # Render page as image
                    mat = fitz.Matrix(dpi/72, dpi/72)  # Convert to desired DPI
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Save image
                    img_path = os.path.join(temp_dir, f"page_{page_num+1}.png")
                    pix.save(img_path)
                    
                    # Add image to Word document
                    doc.add_picture(img_path, width=Inches(7.5))  # Standard page width
                    
                    # Add page break (except for last page)
                    if page_num < len(pdf_document) - 1:
                        doc.add_page_break()
            
            # Save document
            doc.save(output_path)
            pdf_document.close()
            print("Image-based DOCX conversion successful!")
            return True
            
        except Exception as e:
            print(f"Image-based conversion failed: {e}")
            print(traceback.format_exc())
            return False
    
    def convert_to_docx_text_based(self, pdf_path, output_path):
        """Convert PDF to Word with text extraction - FOR TEXT-BASED PDFs"""
        try:
            if not self.libraries['fitz'] or not self.libraries['python_docx']:
                return False
                
            import fitz
            from docx import Document
            
            pdf_document = fitz.open(pdf_path)
            doc = Document()
            
            # Add title
            doc.add_heading(os.path.basename(pdf_path), 0)
            
            # Extract text from each page
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text = page.get_text()
                
                if text.strip():
                    doc.add_heading(f"Page {page_num + 1}", level=1)
                    paragraphs = text.split('\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            doc.add_paragraph(paragraph)
                    
                    if page_num < len(pdf_document) - 1:
                        doc.add_page_break()
                else:
                    # If no text found, this might be a scanned PDF
                    self.add_result(f"⚠ Page {page_num + 1} appears to be scanned - no text found")
            
            doc.save(output_path)
            pdf_document.close()
            print("Text-based DOCX conversion successful!")
            return True
            
        except Exception as e:
            print(f"Text-based conversion failed: {e}")
            return False
    
    def convert_to_image(self, pdf_path, output_path, format):
        """Convert PDF to images"""
        try:
            from pdf2image import convert_from_path
            
            # Create a directory for multiple pages
            base_name = os.path.splitext(output_path)[0]
            
            # Use higher DPI for better quality if option is selected
            dpi = 300 if self.high_quality.get() else 200
            
            images = convert_from_path(pdf_path, dpi=dpi, fmt=format.upper())
            
            if len(images) == 1:
                # Single page - use original output path
                image = images[0]
                # Convert to RGB if necessary (for JPEG)
                if format.upper() == 'JPEG' and image.mode != 'RGB':
                    image = image.convert('RGB')
                image.save(output_path, format=format.upper())
            else:
                # Multiple pages - create directory
                os.makedirs(base_name, exist_ok=True)
                for i, image in enumerate(images):
                    # Convert to RGB if necessary (for JPEG)
                    if format.upper() == 'JPEG' and image.mode != 'RGB':
                        image = image.convert('RGB')
                    page_path = os.path.join(base_name, f"page_{i+1}.{format}")
                    image.save(page_path, format=format.upper())
            
            return True
        except Exception as e:
            print(f"Image conversion error: {e}")
            print(traceback.format_exc())
            return False
    
    def convert_to_text(self, pdf_path, output_path):
        """Convert PDF to text"""
        try:
            if self.libraries['fitz']:
                import fitz
                doc = fitz.open(pdf_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                return True
            elif self.libraries['pypdf']:
                # Fallback to PyPDF2
                import PyPDF2
                with open(pdf_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                return True
            else:
                return False
        except Exception as e:
            print(f"Text conversion error: {e}")
            print(traceback.format_exc())
            return False
    
    def copy_pdf(self, pdf_path, output_path):
        """Copy PDF file (useful for batch processing)"""
        try:
            shutil.copy2(pdf_path, output_path)
            return True
        except Exception as e:
            print(f"PDF copy error: {e}")
            return False
    
    def update_status(self, message):
        self.root.after(0, lambda: self.status_label.config(text=message))
    
    def update_progress(self, value):
        self.root.after(0, lambda: self.progress_bar.config(value=value))
    
    def add_result(self, message):
        self.root.after(0, lambda: self.results_text.insert(tk.END, message + "\n"))
        self.root.after(0, lambda: self.results_text.see(tk.END))
    
    def conversion_complete(self, successful, failed, failed_files):
        self.convert_btn.config(state='normal')
        self.update_status(f"Conversion complete: {successful} successful, {failed} failed")
        
        if failed > 0:
            messagebox.showwarning(
                "Conversion Complete with Errors",
                f"Conversion completed!\n\n"
                f"Successful: {successful}\n"
                f"Failed: {failed}\n\n"
                f"Check the results panel for details."
            )
        else:
            messagebox.showinfo(
                "Conversion Complete",
                f"All files converted successfully!\n\n"
                f"Files saved to: {self.output_path.get()}"
            )

def main():
    root = tk.Tk()
    app = UniversalPDFConverter(root)
    root.mainloop()

if __name__ == "__main__":
    main()
